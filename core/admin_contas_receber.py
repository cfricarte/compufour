from django.contrib import admin
from django.contrib import messages
from django import forms
from django.contrib.admin.helpers import ACTION_CHECKBOX_NAME
from django.conf import settings
from django.template.response import TemplateResponse
from django.urls import reverse
from django.db import transaction
from django.utils.translation import ngettext
from django.db.models import Sum, Q, Value, ExpressionWrapper, Count, Case, When, Max, F
from django.db.models.functions import Coalesce
from django.utils import timezone
from django.utils.html import format_html
from django.http import HttpResponse
from datetime import date
from django.db.models import DecimalField
from decimal import Decimal
from rangefilter.filters import DateRangeFilter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from .models import ContasReceber, Recebimento, Empresa, Cliente, Venda, PlanoConta, Caixa
from rangefilter.filters import DateRangeFilter


class RecebimentoInline(admin.TabularInline):
    model = Recebimento
    extra = 1


class StatusRecebimentoFilter(admin.SimpleListFilter):
    title = 'Status de Recebimento'
    parameter_name = 'status_recebimento'

    def lookups(self, request, model_admin):
        return (
            ('recebida', 'Recebida'),
            ('parcial', 'Recebida Parcialmente'),
            ('pendente_atrasada', 'Pendente (Atrasada)'),
            ('pendente_hoje', 'Pendente (Vencendo Hoje)'),
            ('pendente_a_vencer', 'Pendente (A Vencer)'),
        )

    def queryset(self, request, queryset):
        qs = queryset.annotate(
            total_recebido=Coalesce(
                Sum('recebimento__recebimento_valor_recebido'),
                Value(Decimal('0')),
                output_field=DecimalField(max_digits=10, decimal_places=2)
            )
        )
        today = date.today()

        if self.value() == 'recebida':
            return qs.filter(total_recebido__gte=F('contas_receber_valor'))

        if self.value() == 'parcial':
            return qs.filter(
                total_recebido__gt=0,
                total_recebido__lt=F('contas_receber_valor')
            )

        if self.value() == 'pendente_atrasada':
            return qs.filter(
                Q(total_recebido__lt=F('contas_receber_valor')) | Q(total_recebido__isnull=True),
                contas_receber_data_vencimento__lt=today
            )

        if self.value() == 'pendente_hoje':
            return qs.filter(
                Q(total_recebido__lt=F('contas_receber_valor')) | Q(total_recebido__isnull=True),
                contas_receber_data_vencimento=today
            )

        if self.value() == 'pendente_a_vencer':
            return qs.filter(
                Q(total_recebido__lt=F('contas_receber_valor')) | Q(total_recebido__isnull=True),
                contas_receber_data_vencimento__gt=today
            )

        return queryset


@admin.register(ContasReceber)
class ContasReceberAdmin(admin.ModelAdmin):
    list_display = (
        'empresa',
        'plano_conta',
        'contas_receber_numero_documento',
        #'contas_receber_id',
        'cliente',
        'contas_receber_valor',
        'contas_receber_data_vencimento',
        'total_recebido_display',
        'saldo_devedor_display',
        'exibir_status'
    )
    search_fields = ('cliente__cliente_nome', 'contas_receber_historico')
    list_filter = (
        StatusRecebimentoFilter,
        ('contas_receber_data_vencimento', DateRangeFilter),
        ('contas_receber_data_emissao', DateRangeFilter),
        'empresa',
        'cliente',
        'plano_conta',
    )
    inlines = [RecebimentoInline]
    actions = ['receber_contas_selecionadas', 'gerar_relatorio_word']
    
    # Campos editáveis no formulário
    fields = (
        'empresa',
        'cliente',
        'venda',
        'plano_conta',
        'contas_receber_numero_documento',
        'contas_receber_historico',
        'contas_receber_data_emissao',
        'contas_receber_data_vencimento',
        'contas_receber_valor',
        'contas_receber_portador',
        'contas_receber_nosso_numero',
    )
    
    # Campos obrigatórios
    readonly_fields = ()
    
    # Campos com autocomplete (opcional)
    autocomplete_fields = ['cliente', 'plano_conta']

    _plano_conta_padrao_cache = None

    def _obter_plano_para_recebimento(self, conta):
        if getattr(conta, 'plano_conta_id', None):
            return conta.plano_conta
        if getattr(conta, 'venda', None) and getattr(conta.venda, 'plano_conta_id', None):
            return conta.venda.plano_conta
        if self._plano_conta_padrao_cache is None:
            self._plano_conta_padrao_cache = PlanoConta.objects.filter(pk=1).first()
        return self._plano_conta_padrao_cache

    def _historico_recebimento(self, recebimento):
        conta = recebimento.contas_receber
        numero = conta.contas_receber_numero_documento or conta.contas_receber_id
        return f'Recebimento #{recebimento.pk} da conta {numero}'

    def _registrar_recebimento_no_caixa(self, request, recebimento):
        valor = recebimento.recebimento_valor_recebido or Decimal('0')
        if valor <= Decimal('0'):
            self._remover_recebimento_caixa(recebimento)
            return
        conta = recebimento.contas_receber
        plano = self._obter_plano_para_recebimento(conta)
        if plano is None:
            if request is not None:
                self.message_user(
                    request,
                    'Não foi possível registrar o recebimento no caixa porque o plano de contas não foi definido.',
                    level=messages.WARNING,
                )
            return
        data_recebimento = recebimento.recebimento_data_recebimento or timezone.localdate()
        historico = self._historico_recebimento(recebimento)
        Caixa.objects.update_or_create(
            empresa=conta.empresa,
            caixa_historico=historico,
            defaults={
                'plano_conta': plano,
                'caixa_data_emissao': data_recebimento,
                'caixa_valor_entrada': valor,
                'caixa_valor_saida': Decimal('0'),
            },
        )

    def _remover_recebimento_caixa(self, recebimento):
        conta = recebimento.contas_receber
        if not getattr(conta, 'pk', None):
            return
        historico = self._historico_recebimento(recebimento)
        Caixa.objects.filter(empresa=conta.empresa, caixa_historico=historico).delete()

    def save_formset(self, request, form, formset, change):
        if formset.model is Recebimento:
            instances = formset.save(commit=False)
            for obj in formset.deleted_objects:
                self._remover_recebimento_caixa(obj)
                obj.delete()
            for obj in instances:
                if not obj.recebimento_data_recebimento:
                    obj.recebimento_data_recebimento = timezone.localdate()
                obj.contas_receber = form.instance
                obj.save()
                self._registrar_recebimento_no_caixa(request, obj)
            formset.save_m2m()
        else:
            super().save_formset(request, form, formset, change)

    @admin.action(description='Receber contas selecionadas automaticamente')
    def receber_contas_selecionadas(self, request, queryset):
        queryset = (
            queryset
            .select_related('empresa', 'plano_conta', 'venda', 'venda__plano_conta')
            .annotate(
                total_recebido=Coalesce(
                    Sum('recebimento__recebimento_valor_recebido'),
                    Value(Decimal('0')),
                    output_field=DecimalField(max_digits=10, decimal_places=2)
                )
            )
        )

        if request.POST.get('post'):
            hoje = date.today()
            contas_recebidas = 0
            contas_ja_quitadas = 0
            contas_sem_valor = 0
            contas_sem_plano = 0
            plano_conta_padrao = None

            with transaction.atomic():
                for conta in queryset.select_for_update():
                    valor_total = conta.contas_receber_valor or Decimal('0')
                    if valor_total <= Decimal('0'):
                        contas_sem_valor += 1
                        continue

                    total_recebido = conta.total_recebido or Decimal('0')
                    saldo = valor_total - total_recebido
                    if saldo <= Decimal('0'):
                        contas_ja_quitadas += 1
                        continue

                    if plano_conta_padrao is None:
                        plano_conta_padrao = PlanoConta.objects.filter(pk=1).first()
                    plano_caixa = (
                        conta.plano_conta
                        or (getattr(conta, 'venda', None) and getattr(conta.venda, 'plano_conta', None))
                        or plano_conta_padrao
                    )
                    if plano_caixa is None:
                        contas_sem_plano += 1
                        continue

                    recebimento = Recebimento.objects.create(
                        contas_receber=conta,
                        recebimento_data_recebimento=hoje,
                        recebimento_valor_recebido=saldo,
                        recebimento_forma_recebimento='Automatico (admin)',
                        recebimento_observacao='Recebimento gerado pela acao em massa do admin.',
                    )

                    historico_recebimento = self._historico_recebimento(recebimento)
                    Caixa.objects.update_or_create(
                        empresa=conta.empresa,
                        caixa_historico=historico_recebimento,
                        defaults={
                            'plano_conta': plano_caixa,
                            'caixa_data_emissao': hoje,
                            'caixa_valor_entrada': saldo,
                            'caixa_valor_saida': Decimal('0'),
                        },
                    )
                    contas_recebidas += 1

            if contas_recebidas:
                self.message_user(
                    request,
                    ngettext(
                        '%d conta a receber foi quitada automaticamente.',
                        '%d contas a receber foram quitadas automaticamente.',
                        contas_recebidas,
                    ) % contas_recebidas,
                    messages.SUCCESS,
                )
            if contas_ja_quitadas:
                self.message_user(
                    request,
                    ngettext(
                        '%d conta ja estava quitada e foi ignorada.',
                        '%d contas ja estavam quitadas e foram ignoradas.',
                        contas_ja_quitadas,
                    ) % contas_ja_quitadas,
                    messages.WARNING,
                )
            if contas_sem_valor:
                self.message_user(
                    request,
                    ngettext(
                        '%d conta sem valor definido nao pode ser recebida.',
                        '%d contas sem valor definido nao puderam ser recebidas.',
                        contas_sem_valor,
                    ) % contas_sem_valor,
                    messages.WARNING,
                )
            if contas_sem_plano:
                self.message_user(
                    request,
                    ngettext(
                        '%d conta sem plano de contas definido nao pode ser registrada no caixa.',
                        '%d contas sem plano de contas definido nao puderam ser registradas no caixa.',
                        contas_sem_plano,
                    ) % contas_sem_plano,
                    messages.WARNING,
                )
            if (
                contas_recebidas == 0
                and contas_ja_quitadas == 0
                and contas_sem_valor == 0
                and contas_sem_plano == 0
            ):
                self.message_user(
                    request,
                    'Nenhuma acao foi realizada.',
                    messages.INFO,
                )
            return None

        contas = list(queryset)
        if not contas:
            self.message_user(
                request,
                'Nenhuma conta foi selecionada para recebimento.',
                messages.WARNING,
            )
            return None

        resumo_contas = []
        total_valor = Decimal('0')
        total_recebido = Decimal('0')
        total_saldo = Decimal('0')

        for conta in contas:
            valor_total = conta.contas_receber_valor or Decimal('0')
            valor_recebido = conta.total_recebido or Decimal('0')
            saldo = max(Decimal('0'), valor_total - valor_recebido)
            total_valor += valor_total
            total_recebido += valor_recebido
            total_saldo += saldo
            resumo_contas.append(
                {
                    'obj': conta,
                    'valor_total': valor_total,
                    'valor_recebido': valor_recebido,
                    'saldo': saldo,
                }
            )

        opts = self.model._meta
        cancel_url = reverse(f"admin:{opts.app_label}_{opts.model_name}_changelist")

        context = {
            **self.admin_site.each_context(request),
            'opts': opts,
            'contas': resumo_contas,
            'total_valor': total_valor,
            'total_recebido': total_recebido,
            'total_saldo': total_saldo,
            'action_checkbox_name': ACTION_CHECKBOX_NAME,
            'selected_ids': request.POST.getlist(ACTION_CHECKBOX_NAME),
            'select_across': request.POST.get('select_across'),
            'index': request.POST.get('index'),
            'action_name': 'receber_contas_selecionadas',
            'title': 'Confirmar recebimento das contas selecionadas',
            'cancel_url': cancel_url,
        }
        request.current_app = self.admin_site.name
        return TemplateResponse(
            request,
            'admin/core/contasreceber/confirmar_recebimento.html',
            context,
        )

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        qs = qs.annotate(
            total_recebido=Coalesce(
                Sum('recebimento__recebimento_valor_recebido'),
                Value(Decimal('0')),
                output_field=DecimalField(max_digits=10, decimal_places=2)
            )
        )
        return qs

    def total_recebido_display(self, obj):
        return f"R$ {obj.total_recebido:.2f}"

    total_recebido_display.short_description = 'Total Recebido'
    total_recebido_display.admin_order_field = 'total_recebido'

    def saldo_devedor_display(self, obj):
        saldo = (obj.contas_receber_valor or 0) - obj.total_recebido
        return f"R$ {saldo:.2f}"

    saldo_devedor_display.short_description = 'Saldo Devedor'

    def exibir_status(self, obj):
        valor_total = obj.contas_receber_valor or 0
        total_recebido = obj.total_recebido

        if total_recebido >= valor_total and valor_total > 0:
            return format_html('<span style="color: green; font-weight: bold;">&#x2714; Recebida</span>')

        status_pendente = ''
        if obj.contas_receber_data_vencimento:
            today = date.today()
            if obj.contas_receber_data_vencimento < today:
                status_pendente = format_html('<span style="color: red; font-weight: bold;">&#x26A0; Atrasada</span>')
            elif obj.contas_receber_data_vencimento == today:
                status_pendente = format_html('<span style="color: orange; font-weight: bold;">&#x23F3; Vencendo Hoje</span>')
            else:
                status_pendente = format_html('<span style="color: blue;">&#x1F4C5; A Vencer</span>')

        if total_recebido > 0:
            return format_html(
                '<span style="color: purple; font-weight: bold;">&#x25CF; Parcial</span><br>{}', status_pendente
            )

        return status_pendente if status_pendente else 'Pendente'

    exibir_status.short_description = 'Status'

    @admin.action(description='📝 Gerar Relatório Word')
    def gerar_relatorio_word(self, request, queryset):
        """
        Gera um relatório em Word (.docx) das contas a receber
        Formato: Documento | Cliente | Emissão | Vencimento | Valor
        """
        # Criar documento Word
        doc = Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Título
        titulo = doc.add_heading('RELATÓRIO DE CONTAS A RECEBER', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data do relatório
        hoje = date.today().strftime('%d/%m/%Y')
        p = doc.add_paragraph(f'Data do Relatório: {hoje}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Espaço
        doc.add_paragraph()
        
        # Ordenar por cliente e depois por vencimento
        queryset = queryset.select_related('cliente').order_by('cliente__cliente_nome', 'contas_receber_data_vencimento')
        
        # Agrupar por cliente
        from itertools import groupby
        contas_por_cliente = {}
        for conta in queryset:
            cliente_nome = conta.cliente.cliente_nome if conta.cliente else 'SEM CLIENTE'
            if cliente_nome not in contas_por_cliente:
                contas_por_cliente[cliente_nome] = []
            contas_por_cliente[cliente_nome].append(conta)
        
        # Tabela com 4 colunas (sem coluna Cliente já que está agrupado)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Cabeçalho
        hdr_cells = table.rows[0].cells
        headers = ['Documento', 'Emissão', 'Vencimento', 'Valor']
        
        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = header
            # Negrito no cabeçalho
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Dados agrupados por cliente
        total_geral = Decimal('0')
        
        for cliente_nome, contas in sorted(contas_por_cliente.items()):
            # Linha de separação com nome do cliente
            cliente_row = table.add_row().cells
            cliente_row[0].merge(cliente_row[3])
            cliente_row[0].text = f'CLIENTE: {cliente_nome}'
            
            # Formatação da linha do cliente
            for paragraph in cliente_row[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Adicionar sombreamento
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E7E6E6')
            cliente_row[0]._element.get_or_add_tcPr().append(shading_elm)
            
            # Repetir cabeçalho das colunas para este cliente
            header_row = table.add_row().cells
            headers = ['Documento', 'Emissão', 'Vencimento', 'Valor']
            
            for i, header in enumerate(headers):
                cell = header_row[i]
                cell.text = header
                # Negrito no cabeçalho
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
                        run.font.name = 'Arial'
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Sombreamento no cabeçalho do grupo
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F2F2F2')
                cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Subtotal do cliente
            subtotal_cliente = Decimal('0')
            
            # Contas do cliente
            for conta in contas:
                row_cells = table.add_row().cells
                
                documento = conta.contas_receber_numero_documento or 'S/N'
                emissao = conta.contas_receber_data_emissao.strftime('%d/%m/%Y') if conta.contas_receber_data_emissao else '-'
                vencimento = conta.contas_receber_data_vencimento.strftime('%d/%m/%Y') if conta.contas_receber_data_vencimento else '-'
                valor = f"R$ {conta.contas_receber_valor:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                
                subtotal_cliente += conta.contas_receber_valor
                total_geral += conta.contas_receber_valor
                
                row_cells[0].text = documento
                row_cells[1].text = emissao
                row_cells[2].text = vencimento
                row_cells[3].text = valor
                
                # Formatação e alinhamento das células de dados
                for i in range(4):
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                            run.font.name = 'Arial'
                
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Linha de subtotal do cliente
            subtotal_row = table.add_row().cells
            subtotal_formatado = f"R$ {subtotal_cliente:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            
            subtotal_row[0].text = ''
            subtotal_row[1].text = ''
            subtotal_row[2].text = f'Subtotal {cliente_nome}:'
            subtotal_row[3].text = subtotal_formatado
            
            # Negrito no subtotal
            for i in [2, 3]:
                for paragraph in subtotal_row[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.italic = True
                        run.font.size = Pt(8)
                        run.font.name = 'Arial'
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Linha de total geral
        total_row = table.add_row().cells
        total_formatado = f"R$ {total_geral:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        
        total_row[0].text = ''
        total_row[1].text = ''
        total_row[2].text = 'TOTAL GERAL:'
        total_row[3].text = total_formatado
        
        # Negrito e sombreamento no total geral
        for i in [2, 3]:
            for paragraph in total_row[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Sombreamento no total geral
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        for i in [0, 1, 2, 3]:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D9D9D9')
            total_row[i]._element.get_or_add_tcPr().append(shading_elm)
        
        # Salvar em memória
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Criar resposta HTTP
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="relatorio_contas_receber.docx"'
        
        # Mensagem de sucesso
        self.message_user(
            request,
            f'Relatório Word gerado com {queryset.count()} conta(s) - Total: {total_formatado}',
            messages.SUCCESS
        )
        
        return response